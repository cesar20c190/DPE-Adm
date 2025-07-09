import streamlit as st
from datetime import date
from streamlit.runtime.scriptrunner import RerunException
import pandas as pd
import io


st.set_page_config(page_title="Termo de Declaração", layout="centered")
st.title("📄 Termo de Declaração")

# Dados principais
data_termo = st.date_input(
    "📅 Data do termo",
    value=date.today(),
    format="DD/MM/YYYY"
)

# Dicionário com regionais e cidades
regionais_cidades = {
    "1ª REGIONAL": ["COITÉ", "FEIRA DE SANTANA", "IPIRÁ", "IRARÁ", "RIACHÃO DO JACUÍPE", "SANTO ESTEVÃO", "SERRINHA"],
    "2ª REGIONAL": ["ITAPETINGA", "POÇÕES", "VITÓRIA DA CONQUISTA"],
    "3ª REGIONAL": ["CANAVIEIRAS", "ILHÉUS"],
    "4ª REGIONAL": ["CAMACAN", "ITABUNA"],
    "5ª REGIONAL": ["CAMPO FORMOSO", "JACOBIN", "JUAZEIRO", "SENHOR DO BOMFIN"],
    "6ª REGIONAL": ["AMARGOSA", "CACHOEIRA", "CRUZ DAS ALMAS", "NAZARÉ", "SANTO AMARO", "SANTO ANTÔNIO DE JESUS", "VALENÇA"],
    "7ª REGIONAL": ["CAMAÇARI", "CANDEIAS", "ITAPARICA", "LAURO DE FREITAS", "SIMÕES FILHO"],
    "8ª REGIONAL": ["BARREIRAS", "BOM JESUS DA LAPA", "LUÍS EDUARDO MAGALHÃES", "MACAÚBAS", "SANTA MARIA DA VITÓRIA"],
    "9ª REGIONAL": ["EUNÁPOLIS", "PORTO SEGURO"],
    "10ª REGIONAL": ["EUCLIDES DA CUNHA", "PAULO AFONSO", "PARIPIRANGA", "RIBEIRA DO POMBAL"],
    "11ª REGIONAL": ["IRECÊ", "ITABERABA", "SEABRA"],
    "12ª REGIONAL": ["IPIAÚ", "JEQUIÉ"],
    "13ª REGIONAL": ["ALAGOINHAS", "CATU", "ESPLANADA"],
    "14ª REGIONAL": ["TEIXEIRA DE FREITAS"],
    "15ª REGIONAL": ["BRUMADO", "GUANAMBI"]
}
abas = st.tabs(["Dados iniciais", "Demanda", "Declaração"])

with abas[0]:
    # Lista de regionais
    lista_regionais = list(regionais_cidades.keys())

    # Campo para selecionar a regional
    regional = st.selectbox("Selecione a Regional", lista_regionais)

    # Campo para selecionar a cidade associada
    cidade = st.selectbox("Selecione a Cidade", regionais_cidades[regional])

    # Exibir confirmação
    st.success(f"📝 Termo para a cidade de **{cidade}** na **{regional}**, data: {data_termo.strftime('%d/%m/%Y')}")
    
    # 1. Identificação
    st.header("Identificação")
    Nome = st.text_input("Nome do assistido", placeholder="Ex: João da Silva")
    cpf = st.text_input("CPF do assistido (apenas números)", placeholder="Ex: 123.456.789-00")

    # 2. Dados do assistido
    st.header("Dados do assistido")
    col1, col2, col3 = st.columns([2, 1, 2])
    with col1:
        rua = st.text_input("Rua do assistido", placeholder="Ex: Rua das Flores")
    with col2:
        numero = st.text_input("Número", placeholder="Ex: 123")
    with col3:
        bairro = st.text_input("Bairro", placeholder="Ex: Centro")

    cidade_assistido = st.text_input("Cidade do assistido", placeholder="Ex: Salvador")
    telefone = st.text_input("Telefone do assistido", placeholder="Ex: (71) 99999-9999")
    datanascimento = st.date_input("Data de nascimento do assistido", format="DD/MM/YYYY")


    sexo = st.selectbox("Sexo do assistido", ["Masculino", "Feminino", "Homen trans", "Mulher trans", "Não binário", "Prefiro não informar"])
    grupo_etinico = st.selectbox("Grupo étnico do assistido", ["Branco", "Preto", "Pardo", "Amarelo", "Indígena", "Prefiro não informar"])  

    # 3. Dados da demanda
    st.header("Dados da demanda")
    renda_individual = st.selectbox("Renda individual do assistido", ["Nenhuma", "Até 1 salário mínimo", "De 1 a 2 salários mínimos", "de 2 a 3 salários mínimos", "de 3 a 4 salários mínimos", "de 4 a 5 salários mínimos", "mais de 5 salários mínimos"])
    renda_familiar = st.selectbox("Renda familiar do assistido", ["Nenhuma", "Até 1 salário mínimo", "De 1 a 2 salários mínimos", "de 2 a 3 salários mínimos", "de 3 a 4 salários mínimos", "de 4 a 5 salários mínimos", "mais de 5 salários mínimos"])
    materia = st.selectbox("Matéria do termo", ["Saúde", "Tributário", "Indenizatória", "Concuro Público", "Petitórias/Possessórias", "Alvarás", "Outros"])

with abas[1]:
    # Lista para armazenar múltiplas demandas
    if "demandas" not in st.session_state:
        st.session_state["demandas"] = []

    st.header("Demandas do assistido")

    #    Botão para inserir nova demanda
    if st.button("Inserir demanda"):
        if "mostrar_nova_demanda" not in st.session_state:
            st.session_state["mostrar_nova_demanda"] = 1
        else:
            st.session_state["mostrar_nova_demanda"] += 1

    # Exibe campos para cada demanda a ser cadastrada
    num_demandas = st.session_state.get("mostrar_nova_demanda", 0)
    for i in range(num_demandas):
        with st.expander(f"Nova Demanda {i+1}", expanded=True):
            tipo = st.selectbox(
                f"Tipo de demanda {i+1}",
                [
                    "Cirurgia", "Exames", "Medicamentos", "Tratamento", "Insumos", "Consultas",
                    "Tratamento/Terapias continuadas", "Transferência Hospitalar", "Internação em UTI",
                    "Transporte para fora do domicílio-TFD", "Outros"
                ],
                key=f"tipo_demanda_{i}"
            )
            if tipo == "Outros":
                tipo = st.text_input(f"Especifique o tipo de demanda {i+1}", key=f"outros_tipo_{i}")
            descricao = st.text_area(
                f"Descrição da demanda {i+1}",
                placeholder="Descreva a demanda de forma detalhada",
                height=100,
                key=f"descricao_demanda_{i}"
            )
                    
            # Exibe opções adicionais se o tipo for "Medicamentos"
            if tipo == "Medicamentos":
                tipo_medicamento = st.selectbox(
                    "Não Fornecido pelo SUS ou Fornecido pelo SUS?",
                    ["Não fornecido pelo SUS", "Fornecido pelo SUS"],
                    key=f"tipo_medicamento_{i}"
                )
                tipo = f"Medicamentos - {tipo_medicamento}"

            # Salva automaticamente ao preencher ambos os campos
            if tipo and descricao:
                # Cada demanda é salva como uma linha independente
                if len(st.session_state["demandas"]) < i + 1:
                    st.session_state["demandas"].append({
                        "tipo": tipo,
                        "descricao": descricao
                    })
                else:
                    st.session_state["demandas"][i] = {
                        "tipo": tipo,
                        "descricao": descricao
                    }

    # Exibe tabela das demandas cadastradas, uma por linha
    if st.session_state["demandas"]:
        st.subheader("Demandas cadastradas")
        df_demandas = pd.DataFrame(st.session_state["demandas"])
        st.dataframe(df_demandas, use_container_width=True)

with abas[2]:
    import sqlite3
    from docx import Document
    from datetime import datetime
    import os
    import io

    st.header("Declaração")

    declaracao = st.text_area(
        "Declaração",
        placeholder="Digite a declaração aqui...",
        height=200,
        value=f"Assistido(a) declara que..."
    )

    # Estado da sessão para controlar exibição do botão de gerar termo
    if "registro_salvo" not in st.session_state:
        st.session_state["registro_salvo"] = False

    # Botão SALVAR REGISTRO no banco de dados
if st.button("💾 Salvar Registro no Banco de Dados"):
    data_formatada = str(data_termo)

    conexao = sqlite3.connect("banco_termos.db")
    cursor = conexao.cursor()

    # Inserir o termo principal
    cursor.execute("""
        INSERT INTO termos (
            data_termo, regional, cidade_atendimento,
            nome, cpf, rua, numero, bairro, cidade_assistido, telefone,
            data_nascimento, sexo, grupo_etnico, renda_individual, renda_familiar,
            materia, declaracao
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        data_formatada, regional, cidade,
        Nome, cpf, rua, numero, bairro, cidade_assistido, telefone,
        datanascimento.strftime("%Y-%m-%d"), sexo, grupo_etinico,
        renda_individual, renda_familiar, materia, declaracao
    ))

    id_termo = cursor.lastrowid  # Captura ID do termo salvo

    # Inserir cada demanda
    for demanda in st.session_state["demandas"]:
        cursor.execute("""
            INSERT INTO demandas (id_termo, tipo, descricao)
            VALUES (?, ?, ?)
        """, (id_termo, demanda["tipo"], demanda["descricao"]))

    conexao.commit()
    conexao.close()

    st.success("✅ Registro completo salvo no banco de dados.")
    st.session_state["registro_salvo"] = True

    # Botão de gerar documento só aparece após salvar no banco
    if st.session_state["registro_salvo"]:
        st.markdown("---")
        st.subheader("📄 Gerar Documento")

        if st.button("📄 Gerar Termo de Declaração"):
            caminho_modelo = "Documentos/termo_de_declaracao.docx"
            if not os.path.exists(caminho_modelo):
                st.error(f"❌ Arquivo de modelo não encontrado: {caminho_modelo}")
            else:
                doc = Document(caminho_modelo)

                hora_atendimento = datetime.now().strftime("%H:%M")
                data_formatada = data_termo.strftime("%d/%m/%Y")
                qualificacao = f"{sexo}, CPF {cpf}, residente na {rua}, nº {numero}, bairro {bairro}, cidade de {cidade_assistido}, telefone {telefone}, nascido(a) em {datanascimento.strftime('%d/%m/%Y')}, do grupo étnico {grupo_etinico}"

                for p in doc.paragraphs:
                    p.text = p.text.replace("<<nomeassistido>>", Nome)
                    p.text = p.text.replace("<<horaatendimento>>", data_formatada + " às " + hora_atendimento)
                    p.text = p.text.replace("<<qualificacao>>", qualificacao)
                    p.text = p.text.replace("<<declaracao>>", declaracao)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                nome_arquivo = f"Termo_{Nome.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

                st.success("✅ Documento gerado com sucesso.")
                st.download_button(
                    label="📥 Baixar Termo",
                    data=buffer,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # Botão de gerar documento só aparece após salvar no banco
    if st.session_state["registro_salvo"]:
        st.markdown("---")
        st.subheader("📄 Gerar Documento")

        if st.button("📄 Gerar Termo de Declaração"):
            caminho_modelo = "Documentos/termo_de_declaracao.docx"
            if not os.path.exists(caminho_modelo):
                st.error(f"❌ Arquivo de modelo não encontrado: {caminho_modelo}")
            else:
                doc = Document(caminho_modelo)

                hora_atendimento = datetime.now().strftime("%H:%M")
                data_formatada = data_termo.strftime("%d/%m/%Y")
                qualificacao = f"{sexo}, CPF {cpf}, residente na {rua}, nº {numero}, bairro {bairro}, cidade de {cidade_assistido}, telefone {telefone}, nascido(a) em {datanascimento.strftime('%d/%m/%Y')}, do grupo étnico {grupo_etinico}"

                for p in doc.paragraphs:
                    p.text = p.text.replace("<<nomeassistido>>", Nome)
                    p.text = p.text.replace("<<horaatendimento>>", data_formatada + " às " + hora_atendimento)
                    p.text = p.text.replace("<<qualificacao>>", qualificacao)
                    p.text = p.text.replace("<<declaracao>>", declaracao)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                nome_arquivo = f"Termo_{Nome.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"

                st.success("✅ Documento gerado com sucesso.")
                st.download_button(
                    label="📥 Baixar Termo",
                    data=buffer,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )